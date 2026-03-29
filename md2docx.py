"""md2docx - 右键菜单将 Markdown 转换为 DOCX"""

import subprocess
import sys
import os
import re
import ctypes
import tempfile


def show_error(title, message):
    ctypes.windll.user32.MessageBoxW(0, message, title, 0x10)


def show_info(title, message):
    ctypes.windll.user32.MessageBoxW(0, message, title, 0x40)


def _make_border(val, sz="12", space="0", color="000000"):
    """创建一个边框 XML 元素。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    border = OxmlElement(f"w:{val}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    return border


def _apply_three_line_table(table):
    """将表格设为三线表：顶线、表头底线（粗），底线（粗），其余无边框。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl

    # 清除已有表格边框和样式
    tPr = tbl.tblPr
    if tPr is None:
        tPr = OxmlElement("w:tblPr")
        tbl.insert(0, tPr)

    # 移除已有边框
    for old_borders in tPr.findall(qn("w:tblBorders")):
        tPr.remove(old_borders)

    # 设置表格级边框：只有顶线和底线（粗线 1.5pt = sz 12）
    borders = OxmlElement("w:tblBorders")
    borders.append(_make_border("top", sz="12"))      # 顶线 1.5pt
    borders.append(_make_border("bottom", sz="12"))    # 底线 1.5pt
    # 左右和内部横竖线都设为 none
    for side in ("left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tPr.append(borders)

    # 表头行底部加粗线（第一行）
    if len(table.rows) > 1:
        for cell in table.rows[0].cells:
            tcPr = cell._element.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._element.insert(0, tcPr)
            # 移除已有单元格边框
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tc_borders = OxmlElement("w:tcBorders")
            tc_borders.append(_make_border("bottom", sz="12"))  # 表头底线 1.5pt
            tcPr.append(tc_borders)


def apply_styles(docx_path):
    """后处理 DOCX：统一字体、行距、对齐、颜色，移除水平线。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT_SIZE = Pt(12)  # 小四 = 12pt

    doc = Document(docx_path)

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    # 设置中文字体为宋体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置默认段落格式
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 遍历所有段落
    paragraphs_to_remove = []
    for para in doc.paragraphs:
        # 检测水平线段落（pandoc 用边框模拟 <hr>）
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                # 有边框且文本为空 = 水平线
                if not para.text.strip():
                    paragraphs_to_remove.append(para)
                    continue
                # 非空文本段落也移除边框
                pPr.remove(pBdr)

        # 段落格式
        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 每个 run 的字体
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = FONT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = r.makeelement(qn("w:rPr"), {})
                r.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")

    # 删除水平线段落
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

    # 处理表格：三线表 + 字体样式
    for table in doc.tables:
        _apply_three_line_table(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 2.0
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = FONT_SIZE
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        r = run._element
                        rPr = r.find(qn("w:rPr"))
                        if rPr is None:
                            rPr = r.makeelement(qn("w:rPr"), {})
                            r.insert(0, rPr)
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            rFonts = rPr.makeelement(qn("w:rFonts"), {})
                            rPr.insert(0, rFonts)
                        rFonts.set(qn("w:eastAsia"), "宋体")

    # 同时处理所有内置样式的字体
    for s in doc.styles:
        if hasattr(s, "font") and s.font is not None:
            s.font.name = "Times New Roman"
            s.font.size = FONT_SIZE
            s.font.color.rgb = RGBColor(0, 0, 0)
            if hasattr(s, "element") and hasattr(s.element, "rPr") and s.element.rPr is not None:
                s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if hasattr(s, "paragraph_format") and s.paragraph_format is not None:
            s.paragraph_format.line_spacing = 2.0
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(0)

    doc.save(docx_path)


def _find_image(filename, md_dir):
    """在 md 所在目录及其子目录中查找图片文件，返回相对路径。"""
    # 先检查同目录
    candidate = os.path.join(md_dir, filename)
    if os.path.isfile(candidate):
        return filename
    # 递归搜索子目录
    for root, _dirs, files in os.walk(md_dir):
        if filename in files:
            return os.path.relpath(os.path.join(root, filename), md_dir).replace('\\', '/')
    return filename  # 找不到就���样返回


def _preprocess_obsidian_images(md_text, md_dir):
    """将 Obsidian 图片语法 ![[file|size]] 转换为标准 Markdown 图片语法。"""
    def _replace_img(m):
        content = m.group(1)
        # 分离文件名和尺寸参数（如 图片1.png|671）
        parts = content.split('|', 1)
        filename = parts[0].strip()
        # 在目录中查找实际路径
        rel_path = _find_image(filename, md_dir)
        alt = os.path.splitext(filename)[0]
        return f'![{alt}]({rel_path})'

    # 匹配 ![[xxx]] 和 ![[xxx|size]]
    return re.sub(r'!\[\[([^\]]+)\]\]', _replace_img, md_text)


def _preprocess_footnotes(md_text):
    """将各种引文/脚注语法转为行内文本，避免 pandoc 生成脚注。

    处理的格式：
    - ^[1]  ^[2-4]  ^[note text]   → [1]  [2-4]  [note text]  (pandoc 行内脚注)
    - ^{1-5}  ^{6,7}               → [1-5]  [6,7]            (LaTeX 上标引用)
    - [^id] / [^id]: definition     → [n] + 文末列表          (标准 markdown 脚注)
    """
    # 1) pandoc 行内脚注: ^[...] → [...]
    md_text = re.sub(r'\^\[([^\]]+)\]', r'[\1]', md_text)

    # 2) LaTeX 上标引用: ^{...} → [...]
    md_text = re.sub(r'\^\{([^}]+)\}', r'[\1]', md_text)

    # 3) 标准 markdown 脚注: [^id]: definition 和 [^id] 引用
    footnote_def_re = re.compile(
        r'^\[\^([^\]]+)\]:\s*(.*(?:\n(?![\[\n])(?:[ \t]+.*))*)',
        re.MULTILINE,
    )
    definitions = {}
    for m in footnote_def_re.finditer(md_text):
        fid = m.group(1)
        content = re.sub(r'\n[ \t]+', ' ', m.group(2)).strip()
        definitions[fid] = content

    if not definitions:
        return md_text

    cleaned = footnote_def_re.sub('', md_text)

    order = []
    def _replace_ref(m):
        fid = m.group(1)
        if fid not in order:
            order.append(fid)
        num = order.index(fid) + 1
        return f'[{num}]'

    cleaned = re.sub(r'\[\^([^\]]+)\]', _replace_ref, cleaned)

    # 清理末尾多余空行
    cleaned = cleaned.rstrip('\n') + '\n'

    # 生成参考文献列表
    if order:
        ref_lines = ['\n---\n']
        for i, fid in enumerate(order, 1):
            content = definitions.get(fid, '')
            ref_lines.append(f'[{i}] {content}')
        cleaned += '\n'.join(ref_lines) + '\n'

    return cleaned


def convert(md_path):
    if not os.path.isfile(md_path):
        show_error("md2docx", f"文件不存在：\n{md_path}")
        return

    docx_path = os.path.splitext(md_path)[0] + ".docx"
    md_dir = os.path.dirname(os.path.abspath(md_path))

    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
    except FileNotFoundError:
        show_error("md2docx", "未检测到 pandoc，请先安装：\nhttps://pandoc.org/installing.html")
        return

    # 预处理：脚注转行内引用
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        md_text = _preprocess_obsidian_images(md_text, md_dir)
        md_text = _preprocess_footnotes(md_text)
    except Exception as e:
        show_error("md2docx", f"读取 Markdown 失败：\n{str(e)}")
        return

    # 写入临时文件供 pandoc 使用
    tmp_md = None
    try:
        tmp_fd, tmp_md = tempfile.mkstemp(suffix=".md", dir=md_dir)
        with os.fdopen(tmp_fd, "w", encoding="utf-8") as f:
            f.write(md_text)

        subprocess.run(
            [
                "pandoc", tmp_md, "-o", docx_path,
                "--from", "markdown", "--to", "docx",
                "--resource-path", md_dir,
            ],
            capture_output=True,
            text=True,
            check=True,
            cwd=md_dir,
        )
    except subprocess.CalledProcessError as e:
        show_error("md2docx", f"转换失败：\n{e.stderr}")
        return
    finally:
        if tmp_md and os.path.exists(tmp_md):
            os.unlink(tmp_md)

    try:
        apply_styles(docx_path)
        show_info("md2docx", f"转换成功！\n{docx_path}")
    except ImportError:
        show_error("md2docx", "缺少 python-docx 库，请运行：\npip install python-docx")
    except Exception as e:
        show_error("md2docx", f"样式处理失败：\n{str(e)}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        show_error("md2docx", "用法：md2docx.py <文件路径.md>")
        sys.exit(1)
    convert(sys.argv[1])
